from pathlib import Path
import re
import subprocess
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

base = Path('/Users/evanhan/项目/光伏项目')
word_dir = base / '比赛材料' / 'Word定稿'
total = base / '比赛材料' / 'BP总稿_第一轮精修版.md'
out_md = word_dir / 'SolarGuard商业计划书_装配版_第三版.md'
out_docx = word_dir / 'SolarGuard商业计划书_互联网杯与大创省赛版_第三版.docx'

def build_markdown():
    text = total.read_text(encoding='utf-8')
    start = text.index('# 第一章 项目概述')
    body = text[start:]
    body = re.sub(r'^> 建议.*\n?', '', body, flags=re.M)
    body = re.sub(r'\[当前采用图稿：`([^`]+)`\]', lambda m: f'![插图]({m.group(1)})', body)

    image_map = {
        '合规边界示意图': None,
        '前端原型界面图': str(base / '比赛材料' / '素材库' / '最终图片' / '07_前端页面展示图_v6.png'),
        '应用场景矩阵图': None,
        '市场机会逻辑图': None,
        '商业模式画布简化图': None,
        '推广路径漏斗图': None,
        '项目形成路径图': str(base / '比赛材料' / '素材库' / '最终图片' / '18_项目形成路径图_v1.png'),
        '附录结构导览图': None,
    }
    for label, path in image_map.items():
        pat = f'[图片占位：{label}]'
        if path:
            body = body.replace(pat, f'![{label}]({path})')
        else:
            body = body.replace(pat, f'> 说明：{label} 当前不放入正文图片版，后续可在附录或补充材料中加入。')

    team_page = '''## 7.4 团队能力结构

### 7.4.1 SolarGuard 团队协同结构

以技术实现为基础，以政策研判、产品组织和展示表达为支撑，形成完整参赛产出。

### 7.4.2 技术研发能力
设备接入、自动监测、联盟链交互、前后端原型。
形成结果：构建可运行的系统底座。

### 7.4.3 产品与场景能力
用户对象梳理、功能模块定义、场景延展判断、解决方案表达。
形成结果：把技术能力转为可解释的产品结构。

### 7.4.4 政策与材料能力
政策搜集、合规边界分析、商业计划书撰写、申报材料整合。
形成结果：保证项目叙事和比赛口径稳定。

### 7.4.5 展示与传播能力
图表制作、页面呈现、GitHub 维护、答辩与路演支持。
形成结果：把原型整理为可展示成果。

### 7.4.6 协同结果
代码、实验、图表、政策材料、商业计划书和答辩内容可以被统一组织为一套完整项目资产。

### 7.4.7 核心判断
项目依靠的不是单一点能力，而是多模块协同把复杂问题稳定推进为完整原型。
'''

    edu_page = '''## 7.6 教育实效总结

### 7.6.1 知识迁移
把通信协议、程序设计、数据处理、前后端与区块链知识迁移到真实复杂问题中。

### 7.6.2 工程实践
围绕设备、脚本、接口、页面和系统运行逻辑，形成完整原型，而不是孤立代码片段。

### 7.6.3 研究论证
同步进行政策研判、场景分析、商业论证和合规边界判断。

### 7.6.4 综合表达与组织输出
把代码、实验、图表、仓库、BP 和答辩材料组织成可被评委快速理解的项目成果，完成从“做系统”到“讲清系统”的能力转换。

### 7.6.5 教育实效结论
SolarGuard 的训练价值不只在于“做出一个项目”，而在于让成员完成从知识学习到系统实现、从技术验证到现实表达的完整成长闭环。它符合创新大赛“以赛促学、以创促用、专创融合”的核心要求。
'''
    body = re.sub(r'\[页面占位：团队能力结构页[^\]]*\]', team_page, body)
    body = re.sub(r'\[页面占位：教育实效总结页[^\]]*\]', edu_page, body)
    body = re.sub(r'\[表格占位：([^\]]+)\]', r'> 待补表格：\1', body)
    body = re.sub(r'\n{3,}', '\n\n', body)

    cover = '''# SolarGuard：分布式光伏可信存证与绿色资产服务平台

中国国际大学生创新大赛 / 大学生创新创业训练计划项目

互联网杯与大创省赛商业计划书

项目类型：创意组 / 学生创新项目

项目方向：新能源可信数据、联盟链存证、绿色资产服务

团队名称：待补  
学校：待补  
学院：待补  
指导老师：待补  
项目负责人：待补  
完成时间：2026年3月

\\newpage
'''

    abstract = '''# 执行摘要

SolarGuard 是一个面向分布式光伏场景的可信数据采集、联盟链存证与绿色资产服务平台。项目围绕“新能源场景中的底层数据如何可信获取、可信记录并服务上层应用”这一现实问题展开，形成了从设备接入、自动监测、数字签名、链上存证到前端展示和报告输出的完整原型链路。

与传统光伏监控系统相比，SolarGuard 的核心价值不在于简单展示设备运行状态，而在于把分散、易丢失、难验证的运行数据转化为可追溯、可核验、可服务绿证、ESG 与绿色金融辅助决策的可信数据资产基础。项目当前已经完成真实设备读取、自动监测图表生成、链上交互原型、前端可视化页面和系统级商业计划书材料整合，具备继续向比赛申报、试点展示和后续深化推进的现实基础。

在政策环境方面，项目主动根据最新监管口径完成叙事调整，不再将自身定义为高风险的代币化或公开交易项目，而是明确定位为新能源可信数据与绿色资产服务平台。这一定位更加符合当前数据要素、绿色金融、可信数据空间和分布式光伏规范化建设的政策方向，也使项目具备更强的比赛适配度和现实落地可能性。

在技术上，项目已经完成设备数据采集、自动监测脚本、签名处理、FISCO-BCOS 联盟链交互、前端原型和图表输出等关键模块，证明项目并非停留在概念层，而是具备真实系统雏形。在应用上，项目面向分布式光伏运维、园区能源管理、绿证辅助核验、ESG 披露支持和绿色金融辅助决策等场景，具有较强的延展性。在团队训练上，项目体现了从知识迁移、工程实践到材料组织和综合表达的完整创新训练路径，符合创新大赛和大创项目对学生成长与专创融合的要求。

当前阶段，SolarGuard 已完成商业计划书第一版、核心图片体系、政策资料归档和 Word 定稿结构搭建。后续随着市场数据、用户调研、财务资料和真实实验材料的进一步补充，项目有条件进一步提升为更加完整的省赛和网评提交版本。

\\newpage

# 目录

[TOC]

\\newpage

# 图目录

> 说明：请在 Word 中使用题注与目录功能生成图目录。

\\newpage

# 表目录

> 说明：请在 Word 中使用题注与目录功能生成表目录。

\\newpage
'''
    assembled = cover + abstract + body
    out_md.write_text(assembled, encoding='utf-8')

def add_toc_field(paragraph, figures=False, tables=False):
    p = paragraph._p
    for child in list(p):
        p.remove(child)
    r = OxmlElement('w:r')
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    r.append(fld)
    p.append(r)

    r = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    if figures:
        instr.text = ' TOC \\h \\z \\c "Figure" '
    elif tables:
        instr.text = ' TOC \\h \\z \\c "Table" '
    else:
        instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    r.append(instr)
    p.append(r)

    r = OxmlElement('w:r')
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'separate')
    r.append(fld)
    p.append(r)

    paragraph.add_run('右键更新域以生成目录')

    r = OxmlElement('w:r')
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'end')
    r.append(fld)
    p.append(r)

def style_docx():
    doc = Document(out_docx)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.6)
    sec.bottom_margin = Cm(2.4)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.4)

    style_font_map = {
        'Normal': ('宋体', Pt(12)),
        'Body Text': ('宋体', Pt(12)),
        'First Paragraph': ('宋体', Pt(12)),
        'Block Text': ('宋体', Pt(11)),
        'Compact': ('宋体', Pt(11)),
        'Image Caption': ('宋体', Pt(10.5)),
        'Heading 1': ('Arial', Pt(18)),
        'Heading 2': ('Arial', Pt(15)),
        'Heading 3': ('Arial', Pt(13)),
    }
    for style_name, (font_name, size) in style_font_map.items():
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = font_name
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体' if 'Heading' not in style_name else '等线')
            style.font.size = size
            if style_name == 'Heading 1':
                style.font.bold = True
                style.font.color.rgb = RGBColor(0x2E, 0x6B, 0x4F)
            elif style_name == 'Heading 2':
                style.font.bold = True
                style.font.color.rgb = RGBColor(0x35, 0x5C, 0x8A)
            elif style_name == 'Heading 3':
                style.font.bold = True
                style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    started_body = False
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if i <= 5:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                if i == 0:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
                    run.font.size = Pt(24)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x2E, 0x6B, 0x4F)
                elif i == 1:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(14)
                elif i == 2:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(16)
                    run.font.bold = True
                else:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(12)
            continue

        if text in {'目录','图目录','表目录'}:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if para.style.name == 'Heading 1':
            if text.startswith('第一章') or text.startswith('第二章') or text.startswith('第三章') or text.startswith('第四章') or text.startswith('第五章') or text.startswith('第六章') or text.startswith('第七章') or text.startswith('第八章') or text.startswith('第九章') or text.startswith('第十章'):
                if para.runs:
                    para.runs[0].add_break(WD_BREAK.PAGE)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif text in {'执行摘要','目录','图目录','表目录'}:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if text == '右键更新域以生成目录':
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if text.startswith('说明：') or text.startswith('待补表格：'):
            for run in para.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        if para.style.name in ('Body Text','First Paragraph','Normal'):
            fmt = para.paragraph_format
            fmt.line_spacing = 1.5
            fmt.space_after = Pt(6)
            fmt.first_line_indent = Cm(0.74)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == '右键更新域以生成目录' and i > 0:
            prev = doc.paragraphs[i-1].text.strip()
            if prev == '目录':
                add_toc_field(para, figures=False, tables=False)
            elif prev == '图目录':
                add_toc_field(para, figures=True, tables=False)
            elif prev == '表目录':
                add_toc_field(para, figures=False, tables=True)

    doc.save(out_docx)

build_markdown()
subprocess.run([
    'pandoc', str(out_md), '-o', str(out_docx), '--resource-path', str(base), '-V', 'geometry:margin=2.5cm'
], check=True)
style_docx()
print(out_md)
print(out_docx)
